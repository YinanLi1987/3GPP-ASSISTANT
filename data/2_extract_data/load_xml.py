
from docx import Document
from lxml import etree
def process_xml(docx_path,output_path):
    # Load the DOCX file
    doc=Document(docx_path)
    # Extract the XML content
    xml_content=doc.element.xml
    # Parse the XML content using lxml
    root = etree.fromstring(xml_content)
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # Extract text content from the DOCX file (<w:t> elements)
    text_elements = root.xpath('//w:t', namespaces=nsmap)
    text_content = ' '.join([text.text for text in text_elements])
    
    # Extract insertions (<w:ins>) and deletions (<w:del>)
    insertions = root.xpath('//w:ins/w:t', namespaces=nsmap)
    deletions = root.xpath('//w:del/w:t', namespaces=nsmap)
    
    # Create a new DOCX file
    new_doc = Document()

  # Add the main text content
    new_doc.add_heading('Document Text', level=1)
    text_paragraph = new_doc.add_paragraph(text_content)

    # Add insertions (tracked changes)
    new_doc.add_heading('Insertions (Tracked Changes)', level=2)
    if insertions:
        for ins in insertions:
            new_doc.add_paragraph(f"Inserted: {ins.text}")
    else:
        new_doc.add_paragraph("No insertions found.")

    # Add deletions (tracked changes)
    new_doc.add_heading('Deletions (Tracked Changes)', level=2)
    if deletions:
        for del_ in deletions:
            new_doc.add_paragraph(f"Deleted: {del_.text}")
    else:
        new_doc.add_paragraph("No deletions found.")

    # Save the new DOCX file with the XML content inside
    new_doc.save(output_path)
docx_path = 'data/2_extract_data/raw_data_example/Draft R4-2409310 SUL big CR.docx'
output_path = 'data/2_extract_data/raw_data_example/xml_content_in_docx.docx'
process_xml(docx_path,output_path)
